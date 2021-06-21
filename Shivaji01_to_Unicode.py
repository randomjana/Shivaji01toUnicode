#This program is based on https://github.com/sanskrit-coders/tech_hindi_font_converters. I have only created the python version of the same with few modifications

from docx import Document
from docx.text.run import Run
import re
a="Shivaji01"
array_one= ["iM" , "|","K","#", "k","@","G", "g","\=","C","c","J","j","H","z","T","Z","D", "N","q","t","%","Q","d","n","f", "p", "P", "B","b","m","y","r","l","v","S","Y", "s", "h","p`", "`", "\/", "x","~","&","V", "Ù", "E", "*", "w", "É", "$", "Ë", "W", "\<", ">", "_","AaO", "AÝ" , "Aao", "Aa", "A","eo", "e", "š", "[", "]","?","¸", "."," aa"," a","\\","्a","Ý","aO","O","I","U", "u", "R", "ao", "o", "a", "^M" , "M", "Á","³" , "\´" ,  "््", "-","(",")","+","!","/"]
array_two=[  "Mi", "ऽ",
"ख", "ख्", "क","क्","घ्","ग्","ङ",
"छ","च्","झ्","ज्","ञ्",
"ठ","ट","ढ", "ड", "ण्",
"थ्","त","त्","ध्","द","न्",

"फ", "प", "प्", "भ्", "ब्", "म्",
"य्", "र","ल्","व्","श्","ष्","स्", "ह",

"प्र्", "्र", "्र", "क्ष्", "त्र","ज्ञ",  "द्य", "न्न", "श्र्", "ह्म", "द्ध", "रु", "रू", "क्र", "द्व", "त्त्", 'क्त', 'द्द',
"औ", "औ", "ओ", "आ", "अ","ऐ", "ए", "ई", "इ", "उ", "ऋ", "," , "|" , " ||" ,  " |" ,
"्", "" , "ौ", "ौ", "ै", "ी", "ू", "ु", "ॄ", "ो",  "े", "ा", "ँ" , "ं", ":" ,
"\(" , "\)", "्","र्","ह्य","हृ","ट्ट","ॐ","्र"]

document = Document(FileName) #Enter file name here
for paragraph in document.paragraphs:
    runs=paragraph.runs
    for j in range(0,len(runs)):
        if runs[j].font.name == a:
            idx=0
            #Check and pick the characters in Shivaji01 font for conversion
            for k in array_one :
                
                index=0
                while(index!=-1):
                    runs[j].font.name="Shobhika Regular"
                    runs[j].text=runs[j].text.replace(array_one[idx],array_two[idx])
                    index=runs[j].text.find(array_one[idx])
                idx=idx+1
            position_of_i=runs[j].text.find("i")
            while (position_of_i !=-1):
                runs[j].font.name="Shobhika Regular"
                try:
                    character_next_to_i = runs[j].text[position_of_i+1]
                    character_to_be_replaced="i"+character_next_to_i
                except:
                    runs[j].text=runs[j].text+runs[j+1].text
                    character_next_to_i = runs[j].text[position_of_i+1]
                    character_to_be_replaced="i"+character_next_to_i
                finally:
                    runs[j].text=runs[j].text.replace(character_to_be_replaced,character_next_to_i+"ि")
                    position_of_i=runs[j].text.find("i",position_of_i+1)
    #Correct position of 'i'kara
            position_of_wrong_ee=runs[j].text.find("ि्")
            while(position_of_wrong_ee !=-1):
                runs[j].font.name="Shobhika Regular"
                try :
                    consonent_next_to_wrong_ee=runs[j].text[position_of_wrong_ee+2]
                except:
                    runs[j].text=runs[j].text+runs[j+1].text
                    consonent_next_to_wrong_ee=runs[j].text[position_of_wrong_ee+2]
                finally:
                    character_to_be_replaced="ि्" + consonent_next_to_wrong_ee
                    runs[j].text=runs[j].text.replace(character_to_be_replaced , "्" + consonent_next_to_wrong_ee + "ि" )
                    position_of_wrong_ee = runs[j].text.find("ि्",position_of_wrong_ee+2)

document.save(Filename) #Enter file name here
#Position of arka
set_of_matras = "ािीुूृेैोौंःँॅ"
for paragraph in document.paragraphs:

    pos_reph=re.finditer('र्',paragraph.text)
    pos_reph=[match.start() for match in pos_reph]
   
    for i in pos_reph:
        probable_pos_of_half_r=i-1
        char_at_probable_pos_of_half_r=paragraph.text[probable_pos_of_half_r]

        if (re.search(char_at_probable_pos_of_half_r,set_of_matras)!= None):
            n=0
            while(re.search(char_at_probable_pos_of_half_r,set_of_matras)!= None):
                probable_pos_of_half_r=probable_pos_of_half_r-1
                char_at_probable_pos_of_half_r=paragraph.text[probable_pos_of_half_r]
       
                n=n+1
        
            tempnum=i-probable_pos_of_half_r
            char_to_be_replaced=paragraph.text[probable_pos_of_half_r:i]
       
            new_string='र्' + char_to_be_replaced
            old_string=char_to_be_replaced+'र्'
      
            paragraph.text=paragraph.text.replace(old_string,new_string)
          
        else:
            char_to_be_replaced=paragraph.text[probable_pos_of_half_r]
            new_string='र्' + char_to_be_replaced
            old_string=char_to_be_replaced+'र्'
        
            paragraph.text=paragraph.text.replace(old_string,new_string)
          

    
document.save("Bhagavad_gita.docx")                
            

                    
